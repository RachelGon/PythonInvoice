from datetime import date
from datetime import datetime, timedelta

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# Checks if it was a holiday
def holiday(id):
    # New years
    if datetime.now().month == 1 and id == 'week1':
        return '630'
    # Memorial day
    elif datetime.now().month == 5 and id == 'week4':
        return '630'
    # Christmas
    elif datetime.now().month == 12 and id == 'week4':
        return '630'
    else:
        return '720'


doc = docx.Document()

doc.add_picture('logo.png', width=Inches(2), height=Inches(1.25))
paragraph = doc.paragraphs[-1]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

paragraph = doc.add_paragraph().add_run('Pelican 1 Owner, LLC' +
                                        '					' +
                                        'INVOICE No.: ' + datetime.now().strftime("%m-%d") + '02\n' +
                                        '8156 Fiddler’s Creek PKWY\n'
                                        'Naples, FL 34114\n'
                                        'Date: ' + datetime.now().strftime("%m/%d/%y"))
paragraph.font.name = 'Calibri (Body)'
paragraph.font.size = Pt(14)


now = date(datetime.now().year, datetime.now().month, 1)
tuesday = now + timedelta(days=1 - now.weekday())

if tuesday.month != now.month:
    tuesday += timedelta(days=7)

monday = tuesday + timedelta(days=6)
weeks = {}
i = 0

while tuesday.month == now.month:
    i += 1
    weeks["week{0}".format(i)] = tuesday.strftime("%B %d") + ' thru ' + monday.strftime("%B %d")
    tuesday += timedelta(days=7)
    monday += timedelta(days=7)


table = doc.add_table(rows=1, cols=5)
row = table.rows[0].cells
row[0].text = 'DATE'
row[1].text = 'DESCRIPTION'
row[2].text = 'UOM'
row[3].text = 'QTY'
row[4].text = 'AMOUNT'
total = 0


for id, name in weeks.items():
    print(id)
    row = table.add_row().cells
    row[0].text = name
    row[1].text = 'Clean Mystique Sales Ctr'
    row[2].text = 'LS'
    row[3].text = '1'
    row[4].text = '$' + str(holiday(id)) + '.00'
    total += int(holiday(id))

table.style = 'Table Grid'

totalCell = doc.add_table(rows=1, cols=5)
row = totalCell.rows[0].cells
row[3].text = 'Total Amount Due:'
row[4].text = '$' + f'{total:,}' + '.00'

row[4].style = 'Table Grid'

paragraph = doc.add_paragraph().add_run('Please make all checks payable to:\n' +
                                        'Yoly Cleaning Service, 311 6th Street SE, Naples, FL 34117')
paragraph.font.name = 'Calibri (Body)'
paragraph.font.size = Pt(14)

doc.save('Invoice_Mystique_' + datetime.now().strftime("%m-%d-%y") + '.docx')

